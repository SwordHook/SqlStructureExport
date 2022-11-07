import os
import re
from tkinter import filedialog
import tkinter.messagebox as msgbox
import pymysql
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import tkinter as tk
from tkinter import *

win1 = tk.Tk()  # 常见窗口对象
win1.title('SqlStructureExportHelper')  # 添加窗体名称
win1.geometry("550x500")

tk.Label(win1, text="地址：").grid(row=0, column=0, sticky="e", padx=(150, 50))
tk.Label(win1, text="端口：").grid(row=1, column=0, sticky="e", padx=(150, 50))
tk.Label(win1, text="用户名：").grid(row=2, column=0, sticky="e", padx=(150, 50))
tk.Label(win1, text="密码：").grid(row=3, column=0, sticky="e", padx=(150, 50))

host_entry = tk.Entry(win1)
host_entry.grid(row=0, column=1, pady=10, sticky="w")
port_entry = tk.Entry(win1)
port_entry.grid(row=1, column=1, pady=10, sticky="w")
username_entry = tk.Entry(win1)
username_entry.grid(row=2, column=1, pady=10, sticky="w")
password_entry = tk.Entry(win1)
password_entry.grid(row=3, column=1, pady=10, sticky="w")

host_entry.insert(INSERT, "223.84.77.163")
port_entry.insert(INSERT, 33306)
username_entry.insert(INSERT, "root")
password_entry.insert(INSERT, "leishu2022")

variable = tk.StringVar(win1)

scorllbar = tk.Scrollbar(win1, orient=tk.HORIZONTAL)
listbox1 = Listbox(win1, selectmode=MULTIPLE, xscrollcommand=scorllbar.set)
listbox1.config(width=30)
scorllbar.config(command=listbox1.xview)


def show():
    try:
        mydb = pymysql.connect(
            host=host_entry.get(),  # 指定访问的服务器，本地服务器指定“localhost”，远程服务器指定服务器的ip地址
            port=int(port_entry.get()),  # 指定端口号，范围在0-65535
            user=username_entry.get(),  # 用户名
            password=password_entry.get(),  # 密码：这里一定要注意123456是字符串形式
            # database=e3.get(),  # 数据库的名字
            charset='utf8mb4',  # 数据库的编码方式
        )
    except RuntimeError:
        msgbox.showerror('操作警示', '数据库连接失败')

    cursor = mydb.cursor()
    cursor.execute("show databases")

    a = cursor.fetchall()

    choose_database_option = tk.OptionMenu(win1, variable, *a)

    oMenuWidth = len(str(max(a, key=len))) + 2
    choose_database_option.config(width=oMenuWidth)

    variable.set("点击选择数据库")
    choose_database_option.grid(row=6,column=1, pady=10)

    def callback(*args):
        listbox1.grid(row=7, column=0, pady=10, sticky="e")
        scorllbar.grid(row=8, column=0,padx=(40,0), pady=10, sticky=tk.E + tk.W)
        listbox1.delete(0, END)
        table_name = variable.get()
        remove_chars = '[·’!"\#$%&\'()＃！（）*+,./:;<=>?\@，：?￥★、…．＞【】［］《》？“”‘’\[\\]^`{|}~]+'
        table_name = re.sub(remove_chars, '', table_name)
        table_name = '`' + table_name + '`'
        s = "use {}".format(table_name)
        cursor.execute(s)
        cursor.execute("show tables")
        b = cursor.fetchall()
        for item in b:
            listbox1.insert(END, item)
        bt_export.grid(row=7, columnspan=2, sticky="e", pady=10)

    variable.trace("w", callback)


def exportWord():
    mydb = pymysql.connect(
        host=host_entry.get(),  # 指定访问的服务器，本地服务器指定“localhost”，远程服务器指定服务器的ip地址
        port=int(port_entry.get()),  # 指定端口号，范围在0-65535
        user=username_entry.get(),  # 用户名
        password=password_entry.get(),  # 密码：这里一定要注意123456是字符串形式
        # database=e3.get(),  # 数据库的名字
        charset='utf8mb4',  # 数据库的编码方式
    )

    try:
        Folderpath = filedialog.askdirectory()
        filelist = os.listdir(Folderpath)
        for file in filelist:
            if file.endswith('.xlsx'):
                print(file)
    except FileNotFoundError:
        msgbox.showerror('导出中止', '未传入任何文件夹')
        return

    # 根据表名查询对应的字段相关信息
    def query(tableName):
        remove_chars = '[·’!"\#$%&\'()＃！（）*+,./:;<=>?\@，：?￥★、…．＞【】［］《》？“”‘’\[\\]^`{|}~]+'
        tableName = re.sub(remove_chars, '', tableName)

        table_schema = variable.get()
        table_schema = re.sub(remove_chars, '', table_schema)
        # 打开数据库连接
        cur = mydb.cursor()
        sql = "select b.COLUMN_NAME,b.COLUMN_TYPE,b.IS_NULLABLE,b.COLUMN_COMMENT from (select * from information_schema.`TABLES`  where TABLE_SCHEMA='" + table_schema + "') a right join(select * from information_schema.`COLUMNS` where TABLE_SCHEMA='" + table_schema + "') b on a.TABLE_NAME = b.TABLE_NAME where a.TABLE_NAME='" + str(
            tableName) + "'"
        print(sql)
        cur.execute(sql)
        data = cur.fetchall()
        cur.close
        return data

    # # 查询当前库下面所有的表名，表名：tableName；表名+注释(用于填充至word文档)：concat(TABLE_NAME,'(',TABLE_COMMENT,')')
    # def queryTableName():
    #     cur = mydb.cursor()
    #     sql = "select TABLE_NAME,concat(TABLE_NAME,'(',TABLE_COMMENT,')') from information_schema.`TABLES`  where TABLE_SCHEMA='test_db_test'"
    #     cur.execute(sql)
    #     data = cur.fetchall()
    #     return data

    # 将每个表生成word结构，输出到word文档
    def generateWord(singleTableData, document, tableName):
        p = document.add_paragraph()
        p.paragraph_format.line_spacing = 1.5  # 设置该段落 行间距为 1.5倍
        p.paragraph_format.space_after = Pt(0)  # 设置段落 段后 0 磅
        # document.add_paragraph(tableName,style='ListBullet')
        remove_chars = '[·’!"\#$%&\'()＃！（）*+,./:;<=>?\@，：?￥★、…．＞【】［］《》？“”‘’\[\\]^`{|}~]+'
        tableName = re.sub(remove_chars, '', tableName)
        r = p.add_run('\n' + tableName)
        r.font.name = u'宋体'
        r.font.size = Pt(12)
        table = document.add_table(rows=len(singleTableData) + 1, cols=4, style='Table Grid')
        table.style.font.size = Pt(11)
        table.style.font.name = u'Calibri'
        # 设置表头样式
        # 这里只生成了三个表头，可通过实际需求进行修改
        for i in ((0, '字段名称'), (1, '数据类型'), (2, '是否必填'), (3, '字段注释')):
            run = table.cell(0, i[0]).paragraphs[0].add_run(i[1])
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
            r = run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        for i in range(len(singleTableData)):
            # 设置表格内数据的样式
            for j in range(len(singleTableData[i])):
                run = table.cell(i + 1, j).paragraphs[0].add_run(singleTableData[i][j])
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
                r = run._element
                r.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 定义一个document
    document = Document()
    # 设置字体默认样式
    document.styles['Normal'].font.name = u'宋体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    # 获取当前库下所有的表名信息和表注释信息

    ichose = []
    items = listbox1.curselection()
    if len(items) == 0:
        msgbox.showerror('导出中止', '未选择任何表')
        return

    for i in range(len(items)):
        ichose.append(listbox1.get(items[i]))
    # 循环查询数据库，获取表字段详细信息，并调用generateWord，生成word数据
    # 由于时间匆忙，我这边选择的是直接查询数据库，执行了100多次查询，可以进行优化，查询出所有的表结构，在代码里面将每个表结构进行拆分
    for singleTableName in ichose:
        data = query(str(singleTableName))
        generateWord(data, document, str(singleTableName))
    # 保存至文档
    document.save(Folderpath + '\\数据库设计.docx')
    if len(Folderpath) != 0:
        msgbox.showinfo('导出成功', '数据库结构导出成功')


bt_get_table = tk.Button(win1, text="获取库信息", width=10, command=show)
bt_get_table.grid(row=4, column=1, sticky="w", pady=5)
bt_export = tk.Button(win1, text="选择文件夹导出", width=15, command=exportWord)

win1.mainloop()  # 执行窗体
